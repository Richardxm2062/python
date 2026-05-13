import os
import re
import subprocess
import fitz

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pypdf import PdfReader, PdfWriter

SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"


# =========================================================
# 工具
# =========================================================

def parse_input_paths(s):

    parts = s.split("'")

    paths = [p.strip() for i, p in enumerate(parts) if i % 2 == 1]

    if not paths:
        paths = [line.strip() for line in s.splitlines() if line.strip()]

    return paths


def natural_key(s):

    return [
        int(text) if text.isdigit() else text.lower()
        for text in re.split(r'(\d+)', s)
    ]


# =========================================================
# 提取章节名 & 类型名
# =========================================================

def extract_names(root_dir):

    parent = os.path.basename(os.path.dirname(root_dir))
    current = os.path.basename(root_dir)

    parent_name = re.sub(r'^\d+[-\d]*\s*', '', parent)
    current_name = re.sub(r'^\d+\s*', '', current)

    return parent_name, current_name


# =========================================================
# 页码工具
# =========================================================

def clear_page_numbers(doc):

    for section in doc.sections:

        # =================================================
        # footer
        # =================================================

        section.footer.is_linked_to_previous = False

        footer = section.footer

        footer_element = footer._element

        for child in list(footer_element):

            footer_element.remove(child)

        footer.add_paragraph()

        # =================================================
        # header
        # =================================================

        section.header.is_linked_to_previous = False

        header = section.header

        header_element = header._element

        for child in list(header_element):

            header_element.remove(child)

        header.add_paragraph()


def add_word_page_numbers(doc):

    for section in doc.sections:

        section.footer.is_linked_to_previous = False

        footer = section.footer

        footer_element = footer._element

        for child in list(footer_element):

            footer_element.remove(child)

        p = footer.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


# =========================================================
# docx处理
# =========================================================

def process_docx_format(doc_path):

    # =====================================================
    # 第一阶段：生成无页码PDF
    # =====================================================

    doc_no_page = Document(doc_path)

    clear_page_numbers(doc_no_page)

    temp_docx = doc_path.replace(".docx", "(无页码).docx")

    doc_no_page.save(temp_docx)

    convert_to_pdf(temp_docx)

    os.remove(temp_docx)

    # =====================================================
    # 第二阶段：生成正式页码PDF
    # =====================================================

    doc_with_page = Document(doc_path)

    clear_page_numbers(doc_with_page)

    add_word_page_numbers(doc_with_page)

    doc_with_page.save(doc_path)

    convert_to_pdf(doc_path)


def convert_to_pdf(doc_path):

    subprocess.run([
        SOFFICE,
        "--headless",
        "--convert-to", "pdf",
        doc_path,
        "--outdir", os.path.dirname(doc_path)
    ],
    stdout=subprocess.DEVNULL,
    stderr=subprocess.DEVNULL)


def process_docx(root_dir):

    for root, dirs, files in os.walk(root_dir):

        # =================================================
        # 跳过 word 目录
        # =================================================

        if os.path.basename(root) == "word":
            continue

        word_files = []

        for file in files:

            if file.startswith("~$"):
                continue

            if not file.endswith(".docx"):
                continue

            if "(无页码)" in file:
                continue

            file_path = os.path.join(root, file)

            print(f"[处理docx] {file_path}")

            try:

                process_docx_format(file_path)

                word_files.append(file_path)

            except Exception as e:

                print(f"❌ docx失败: {file_path}")
                print(e)

        # =================================================
        # 移动word
        # =================================================

        if word_files:

            word_dir = os.path.join(root, "word")

            if not os.path.exists(word_dir):
                os.makedirs(word_dir)

            for file_path in word_files:

                filename = os.path.basename(file_path)

                target_path = os.path.join(word_dir, filename)

                if not os.path.exists(target_path):

                    os.rename(file_path, target_path)

                else:

                    print(f"⚠️ 已存在，跳过: {filename}")


# =========================================================
# PDF分类
# =========================================================

def collect_pdfs(root_dir):

    A = []
    B = []

    for root, dirs, files in os.walk(root_dir):

        if os.path.basename(root) == "word":
            continue

        for file in files:

            if not file.endswith(".pdf"):
                continue

            # =================================================
            # 只收集无页码pdf
            # =================================================

            if "(无页码)" not in file:
                continue

            full_path = os.path.join(root, file)

            if "解析" in file:

                A.append(full_path)

            elif "原卷" in file:

                B.append(full_path)

    return A, B


# =========================================================
# 连续页码
# =========================================================

def add_page_numbers(pdf_path):

    doc = fitz.open(pdf_path)

    for i, page in enumerate(doc):

        rect = page.rect

        page.insert_text(

            (rect.width - 40, rect.height - 20),

            str(i + 1),

            fontsize=10,

            fontname="helv",

            color=(0, 0, 0),

            overlay=True
        )

    temp_output = pdf_path.replace(".pdf", "_tmp.pdf")

    doc.save(temp_output)

    doc.close()

    os.remove(pdf_path)

    os.rename(temp_output, pdf_path)


# =========================================================
# 删除临时pdf
# =========================================================

def delete_temp_pdfs(file_list):

    for f in file_list:

        try:

            os.remove(f)

            print(f"[删除临时pdf] {os.path.basename(f)}")

        except:

            pass


# =========================================================
# PDF合并
# =========================================================

def merge_pdfs(file_list, output_path):

    writer = PdfWriter()

    current_page = 0

    for f in file_list:

        name = os.path.basename(f)

        print(f"[合并] {name}")

        try:

            reader = PdfReader(f)

        except:

            print(f"❌ 跳过损坏文件: {f}")

            continue

        num_pages = len(reader.pages)

        # =================================================
        # 书签
        # =================================================

        title = name.replace(".pdf", "")
        title = title.replace("(无页码)", "")
        title = title.replace("（解析版）", "")
        title = title.replace("（原卷版）", "")

        writer.add_outline_item(title, current_page)

        for page in reader.pages:

            writer.add_page(page)

        current_page += num_pages

    with open(output_path, "wb") as f:

        writer.write(f)

    # =====================================================
    # 添加连续页码
    # =====================================================

    add_page_numbers(output_path)

    # =====================================================
    # 删除临时pdf
    # =====================================================

    delete_temp_pdfs(file_list)


# =========================================================
# 主流程
# =========================================================

def process_one(root_dir):

    if not os.path.exists(root_dir):

        print(f"❌ 路径不存在: {root_dir}")

        return

    print("\n==============================")
    print(f"处理路径: {root_dir}")
    print("==============================")

    # =====================================================
    # docx
    # =====================================================

    process_docx(root_dir)

    # =====================================================
    # 收集pdf
    # =====================================================

    A, B = collect_pdfs(root_dir)

    A.sort(key=lambda x: natural_key(os.path.basename(x)))
    B.sort(key=lambda x: natural_key(os.path.basename(x)))

    print(f"解析数量: {len(A)}")
    print(f"原卷数量: {len(B)}")

    parent_name, current_name = extract_names(root_dir)

    # =====================================================
    # 合并
    # =====================================================

    if A:

        filename_A = f"{parent_name}({current_name})(解析).pdf"

        output_A = os.path.join(root_dir, filename_A)

        merge_pdfs(A, output_A)

    if B:

        filename_B = f"{parent_name}({current_name})(原卷).pdf"

        output_B = os.path.join(root_dir, filename_B)

        merge_pdfs(B, output_B)

    print("✅ 完成")


# =========================================================
# 主入口
# =========================================================

def main():

    print("请输入一个或多个路径：")

    user_input = input()

    paths = parse_input_paths(user_input)

    if not paths:

        print("❌ 未识别到路径")

        return

    for p in paths:

        if not os.path.exists(p):

            print(f"❌ 路径不存在: {p}")

            continue

        process_one(p)


if __name__ == "__main__":

    main()