import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from PIL import Image

img_folder = r"/Users/richard/Documents/Physics/高中物理/1 资料/2 专题/10恒定电流/电学实验真题"
output_path = "/Users/richard/Desktop/电学实验讲义.docx"

doc = Document()

# 页面设置
section = doc.sections[0]
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = 914400

page_width = section.page_width
page_height = section.page_height

available_width = page_width - section.left_margin - section.right_margin
available_height = page_height - section.top_margin - section.bottom_margin

SAFETY = 300000
extra = 1000000  # 标题+间距

# ========= 获取“横向铺满”的高度 =========
def get_full_width_height(img_path):
    img = Image.open(img_path)
    w, h = img.size
    scale = available_width / w
    return h * scale

# ========= 插入图片 =========
def add_image(img_path, i, force_shrink=False):
    img = Image.open(img_path)
    w, h = img.size

    scale = available_width / w
    new_h = h * scale

    # 👉 只有在“单图放不下”才缩小
    if force_shrink and new_h > available_height:
        scale = available_height / new_h * scale

    new_w = w * scale

    # 图片
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(img_path, width=new_w)

    # 标题
    raw = os.path.splitext(os.path.basename(img_path))[0]
    clean = raw.replace("_", " ")
    title = f"图{i}：{clean}"

    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(4)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = cap.add_run(title)
    run.bold = True
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(11)

# ========= 数据 =========
files = [os.path.join(img_folder, f) for f in sorted(os.listdir(img_folder)) if f.lower().endswith(".png")]
heights = {f: get_full_width_height(f) for f in files}

unused = files.copy()

i = 1

# ========= 主循环 =========
while unused:
    A = unused.pop(0)
    hA = heights[A]

    best_B = None
    best_fill = 0

    # 👉 尝试找第二张（不缩放前提下）
    for B in unused:
        hB = heights[B]
        total = hA + hB + 2 * extra

        if total <= available_height - SAFETY:
            if total > best_fill:
                best_fill = total
                best_B = B

    # ===== 插入 A =====
    if best_B:
        add_image(A, i, force_shrink=False)
    else:
        # 单图模式 → 允许缩放
        add_image(A, i, force_shrink=True)

    i += 1

    # ===== 插入 B =====
    if best_B:
        add_image(best_B, i, force_shrink=False)
        i += 1
        unused.remove(best_B)

    # 分页
    if unused:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

doc.save(output_path)

print(f"已生成：{output_path}")